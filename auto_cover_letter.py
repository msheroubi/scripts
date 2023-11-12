from docx import Document
from docx.shared import Inches
import pandas as pd

header = "To the Recruitment Team at {company},"

generic_intro =  "I am excited to be applying for the {position} position at {company}. "
generic_outro = "Thank you for your time and consideration. I am excited to learn more about {company}. I look forward to the opportunity to demonstrate my expertise while developing my skills and tackling new challenges. I know how challenging the hiring process can be, please let me know if there is anything I can do to make it easier. "

cert_i = 'I graduated from the University of British Columbia with a BSc. in Computer Science. '

python_i = 'I have been programming in python for the last 6 years, and profesionally for the last 3. '
python_b = {'automation': 'I have used Python for countless automation tasks. ',
            'pandas': 'I have used Python to clean and format data extensively using Pandas. ',
            'webscrape': 'I have used Python for web-scraping using BeautifulSoup and Selenium. ',
            'webdev' : 'I have used Python for web-development with Django. ',
            'etl':'I have used Python to Extract, Transform and Load data from various 3rd party sources (Stripe, Hubspot, Intercom, MongoDB) into a Redshift warehouse. ',
            'neuralnet':'I used Python for image classification with Neural Network applications using TensorFlow and NumPy. '}

data_analysis_i = 'My passion for data combined with my proficiency in Python and SQL would make me a great addition to your team. '
data_analysis_b = 'During my time at Clearmind, I supported our sales, marketing and executive teams with ad hoc analysis using Excel and Python. '

security_i = 'In April 2020, I passed the CompTIA Security+ exam. '
security_b = 'At Clearmind, I took initiative routinely backing up crucial operations data, which helped quickly resolve multiple instances of data loss. Additionally, I was responsible for ensuring our system and policies complied with Canadian Privacy Laws and the EU GDPR. '

startup_i = 'My passion for learning combined with flexibility and diligence would make me a great addition to your startup. '
startup_b = 'I am not a stranger to wearing many hats. At Clearmind, I routinely took on tasks from different departments as needed. '


dbwh_design_b = 'I designed and deployed a data warehouse solution using AWS Redshift for Trellis. The warehouse was designed to house data from multiple sources including Stripe and MongoDB. '

etl_b = "While with Trellis, I designed an ETL solution from the ground-up in Python to migrate data from their production MongoDB. At Clearmind, I built a basic ETL solution using Python to Extract and Transform over 7 years of data to be loaded into a cloud-based solution. "

nlp_i = 'I\'ve always been fascinated by language. From a young age, I would repeat phrases from different languages. As time passed, my gaze shifted to how we understand language and how we can get AI to understand it as well. Natural language processing is a topic I am very passionate about, and would love to pursue a career in this field long-term. '

sql_b = 'While working with Trellis, I designed complex SQL queries to answer key business questions as stored procedures. I continued to keep my SQL skills sharp by completing related questions on websites like leetcode after maxing out the SQL category on hackerrank. '

gen_b = 'At Clearmind, I thoroughly documented the student registration processes and modelled the flow of data through our system. I also helped onboard and train new employees on the company\'s backend tools. '

safespace = 'In both my personal life and professional, I strive to create frictionless relationships. I try my best to make everyone around me feel safe and comfortable. '
growth = "I am constantly looking for ways to grow and learn."

df = pd.read_excel('jobs.xlsx', engine='openpyxl')

for idx, row in df.iterrows():



    position = row[0]
    company = row[1]
    job_desc = row[2]

    if company == '':
        break
    company_abbrv = company.split()[0]



    document = Document()
    x = document.add_paragraph('Michael Sheroubi\n+17788580920\nmsheroubi@gmail.com')

    head = document.add_paragraph(header.format(company=company))

    i = document.add_paragraph(generic_intro.format(position=position,company=company))

    i.add_run(cert_i)

    job_desc = job_desc.lower()

    if "nlp" in job_desc:
        i.add_run(nlp_i)

    if 'analysis' in job_desc:
        i.add_run(data_analysis_i)

    if 'python' in job_desc:
        i.add_run(python_i)

    if 'startup' in job_desc or 'venture' in job_desc:
        i.add_run(startup_i)

    b = document.add_paragraph(gen_b)

    count = 1

    if 'sql' in job_desc:
        b.add_run(sql_b)
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'database' in job_desc:
        b.add_run(dbwh_design_b)

        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'etl' in job_desc:
        b.add_run(etl_b)
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'security' in job_desc:
        b.add_run(security_b)
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'analysis' in job_desc:
        b.add_run(data_analysis_b)
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph() 

    if 'automation' in job_desc or 'script' in job_desc:
        b.add_run(python_b['automation'])
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'scraping' in job_desc:
        b.add_run(python_b['webscrape'])
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'django' in job_desc:
        b.add_run(python_b['webdev'])
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'equality' in job_desc:
        b.add_run(safespace)
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()

    if 'growth' in job_desc:
        b.add_run(growth)
        count+=1
        if count == 3:
            count=0
            b=document.add_paragraph()


    o = document.add_paragraph(generic_outro.format(company=company))

    sig = document.add_paragraph("Looking forward to hearing from you.\nMichael")

    document.save(f"{position} at {company}.docx")
